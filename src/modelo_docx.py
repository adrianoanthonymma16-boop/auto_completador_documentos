"""
Módulo para ler e escrever arquivos DOCX (Microsoft Word)
"""

import re
import os

# Tenta importar python-docx
try:
    from docx import Document
    DOCX_SUPORTADO = True
except ImportError:
    DOCX_SUPORTADO = False


def extrair_placeholders_docx(caminho_docx):
    """
    Extrai todos os placeholders {{...}} de um arquivo DOCX
    
    Args:
        caminho_docx (str): Caminho do arquivo DOCX
    
    Returns:
        set: Conjunto de placeholders encontrados
    """
    if not DOCX_SUPORTADO:
        raise Exception("Suporte a DOCX não disponível. Instale: pip install python-docx")
    
    placeholders = set()
    doc = Document(caminho_docx)
    
    # Verifica parágrafos
    for paragraph in doc.paragraphs:
        matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
        for match in matches:
            placeholders.add(match.strip())
    
    # Verifica tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                for match in matches:
                    placeholders.add(match.strip())
    
    return placeholders


def gerar_docx_preenchido(caminho_modelo, dados, caminho_saida):
    """
    Gera um novo DOCX com os placeholders substituídos
    
    Args:
        caminho_modelo (str): Caminho do arquivo modelo DOCX
        dados (dict): Dicionário {placeholder: valor}
        caminho_saida (str): Onde salvar o novo arquivo
    """
    if not DOCX_SUPORTADO:
        raise Exception("Suporte a DOCX não disponível. Instale: pip install python-docx")
    
    doc = Document(caminho_modelo)
    
    # Substitui em parágrafos
    for paragraph in doc.paragraphs:
        for placeholder, valor in dados.items():
            if f"{{{{{placeholder}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{placeholder}}}}}", valor)
    
    # Substitui em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, valor in dados.items():
                    if f"{{{{{placeholder}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{placeholder}}}}}", valor)
    
    doc.save(caminho_saida)


def docx_suportado():
    """Retorna True se suporte a DOCX está disponível"""
    return DOCX_SUPORTADO